import streamlit as st
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from io import BytesIO
from openai import OpenAI
from dotenv import load_dotenv
import os

def load_docx(file):
    """Load a docx file into a Document object"""
    return Document(file)

def find_section(doc, section_title):
    """Find the paragraph index where a section begins"""
    for i, para in enumerate(doc.paragraphs):
        if section_title.upper() in para.text.strip().upper():
            return i
    return -1

def parse_and_format_bullet(bullet_text):
    """
    Parse bullet text with ** markdown for bold and return segments with formatting info
    
    Example input: "Implemented **CI/CD pipelines** that increased deployment speed by **80%**"
    Example output: [
        ("Implemented ", False),
        ("CI/CD pipelines", True),
        (" that increased deployment speed by ", False),
        ("80%", True)
    ]
    """
    segments = []
    parts = bullet_text.split("**")
    
    # If odd number of ** markers, it's properly formatted
    # If even number, it's malformed (odd number of ** symbols)
    if len(parts) % 2 == 0:
        # Just return the text as-is without bold
        return [(bullet_text, False)]
    
    for i, part in enumerate(parts):
        if part:  # Skip empty strings that can occur with back-to-back **
            # Even-indexed parts are normal text, odd-indexed are bold
            is_bold = (i % 2 == 1)
            segments.append((part, is_bold))
    
    return segments

def add_content_to_section(doc, section_title, content_json):
    """Add all content from JSON to a specific section"""
    section_idx = find_section(doc, section_title)
    
    if section_idx == -1:
        st.error(f"Could not find section: {section_title}")
        return False
    
    # Add JSON content after the section title
    current_idx = section_idx + 1
    
    # For each role/project in the JSON
    for item in content_json:
        for title, bullets in item.items():
            # Add the title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            
            # Format the title
            title_run.font.name = "Times New Roman"
            title_run.font.size = Pt(12)
            title_run.font.italic = True
            title_run.bold = True
            
            # Add space before paragraph for title
            title_para.paragraph_format.space_after = Pt(0)
            
            # Insert at current position
            p = title_para._element
            doc.paragraphs[current_idx]._p.addprevious(p)
            current_idx += 1
            
            # Add bullet points with proper formatting
            for bullet in bullets:
                # Create a new paragraph
                bullet_para = doc.add_paragraph()
                
                
                # Parse the bullet text for bold segments
                segments = parse_and_format_bullet(bullet)
                
                # Add the bullet character first (without making it bold)
                bullet_run = bullet_para.add_run("• ")
                bullet_run.font.name = "Times New Roman"
                bullet_run.font.size = Pt(11)
                
                # Add each text segment with appropriate formatting
                for text, is_bold in segments:
                    run = bullet_para.add_run(text)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run.bold = is_bold
                
                # Set line spacing to 1.0 (single spacing)
                paragraph_format = bullet_para.paragraph_format
                paragraph_format.line_spacing = 1.0
                
                # Remove space after paragraph
                paragraph_format.space_after = Pt(0)
                
                
                
                # Insert at current position
                p = bullet_para._element
                doc.paragraphs[current_idx]._p.addprevious(p)
                current_idx += 1
            if item != content_json[-1] :
              space_para = doc.add_paragraph()
              space_para.paragraph_format.space_after = Pt(0.5)
              p = space_para._element
              doc.paragraphs[current_idx]._p.addprevious(p)
              current_idx += 1 
    return True

def save_docx(doc):
    """Save document to BytesIO object for downloading"""
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def parse_and_format_skill(skill_text):
    """
    Parse skill text with ** or : to determine what should be bold
    
    Example input: "**Programming Languages:** Java, Python, C/C++, JavaScript"
    or: "Programming Languages: Java, Python, C/C++, JavaScript"
    """
    # If already has ** markers, use the standard parsing function
    if "**" in skill_text:
        return parse_and_format_bullet(skill_text)
    
    # If it has a colon, make the part before the colon bold
    elif ":" in skill_text:
        parts = skill_text.split(":", 1)  # Split only at the first colon
        if len(parts) == 2:
            return [(parts[0] + ":", True), (parts[1], False)]
    
    # Otherwise, return as is (not bold)
    return [(skill_text, False)]

def add_skills_to_section(doc, section_title, skills_list):
    """Add skills to the TECHNICAL SKILLS section of the resume"""
    section_idx = find_section(doc, section_title)
    
    if section_idx == -1:
        st.error(f"Could not find section: {section_title}")
        return False
    
    # Add skills after the section title
    current_idx = section_idx + 1
    
    # Process each skill category (e.g., "Programming Languages:", "Software Development:", etc.)
    for skill_category in skills_list:
        # Create a new paragraph for the skill category
        skill_para = doc.add_paragraph()
        
        # Parse the skill text for bold sections (usually the category name)
        segments = parse_and_format_skill(skill_category)
        
        # Add each segment with appropriate formatting
        for text, is_bold in segments:
            run = skill_para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = is_bold
        
        # No spaces before or after paragraph
        paragraph_format = skill_para.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0
        
        # Insert at current position
        p = skill_para._element
        doc.paragraphs[current_idx]._p.addprevious(p)
        current_idx += 1
    
    return True

def process_job_description(resume_json, job_description, model="gpt-4.1-nano"):
    """Process job description with OpenAI API and get enhanced resume content"""
    load_dotenv()
    keymain = os.getenv("API_KEY")
    client = OpenAI(api_key=keymain)
    
    system_prompt = """
    You are a resume optimization assistant. Your task is to extract relevant keywords, skills, technologies, and role responsibilities 
    from the job description. Enhance the candidate's resume bullet points to closely align with these requirements. Prioritize quantifiable results, 
    action verbs, and technical keywords to maximize ATS (Applicant Tracking System) compatibility. Return output in valid JSON format only.
    
    For each bullet point, you can indicate words or phrases that should be bold by surrounding them 
    with **double asterisks**. Use this to highlight key skills, technologies, metrics, or achievements that directly 
    match the job description to draw attention to them. But keep the bolding minimal to one or two in each bullet point. Do not bold more than that.
    
    Example:
    "Implemented **CI/CD pipelines** that increased deployment speed by **80%** while maintaining quality"
    
    For the skills section, make sure to analyze existing skills categories and add relevant skills from the job description to the appropriate category.
    Don't duplicate existing skills. If new skills don't fit into existing categories, add them to the "Additional Skills" category.
    Make sure each skill category header is in bold with format: "**Category Name:** skill1, skill2, skill3"
    """
    
    user_prompt = f"""
    JOB DESCRIPTION:
    {job_description}
    
    RESUME DATA TO ENHANCE:
    {json.dumps(resume_json)}
    
    INSTRUCTIONS:
    1. Carefully extract key responsibilities, required skills, technologies, and qualifications from the job description.
    
    2. Modify each resume bullet point to:
       - Use relevant keywords and phrases from the job description.
       - Include strong action verbs and measurable outcomes in each point (use numbers, percentages, timeframes, etc.).
       - Insert or replace with relevant tools, frameworks, platforms, and languages mentioned in the job posting—or those that are commonly used in such roles.
       - Mark important skills, technologies, and metrics with **bold** (using **double asterisks**) that directly match the job requirements.
       
    3. Add **1–2 new bullet points per section** if necessary, especially to include critical technologies or achievements that would significantly increase ATS relevance.
    
    4. Keep each bullet point concise (1–2 lines max), focused on impact and results.
    
    5. For the "skills" section:
       - Analyze the existing skill categories.
       - Add any new skills from the job description to the appropriate category.
       - Do NOT duplicate skills that already exist.
       - If new skills don't fit existing categories, add them to "Additional Skills".
       - Keep the category header format: "**Category Name:** skill1, skill2, skill3"
       
    6. Output updated content as JSON, keeping the structure identical to the input. Return only valid JSON in this format: 
       {{
         "experience": [
           {{"role_name": ["bullet_point1 with **bold** terms", "bullet_point2", ...]}},
           ...
         ],
         "projects": [
           {{"project_name": ["bullet_point1", "bullet_point2 with **metrics**", ...]}},
           ...
         ],
         "skills": [
           "**Programming Languages:** Java, Python, etc.",
           "**Software Development:** skill1, skill2, etc.",
           ...
         ]
       }}
       
    Goal: Maximize ATS score by aligning resume content with the job description while preserving professional tone and formatting.
    """
    
    try:
        response = client.chat.completions.create(
            model=model,
            response_format={"type": "json_object"},  # Force JSON response
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3  # Lower temperature for more consistent outputs
        )
        
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"Error calling OpenAI API: {str(e)}")
        return None

# Streamlit App
def main():
    st.title("📝 Resume Builder")

    # Sidebar for API key
    model_option = "gpt-4.1-nano"

    # File uploader for resume template
    uploaded_file = "new_resume_half.docx"
    software_dev_resume = {
  "experience": [
    {
      "Software Developer | University of Florida | Nov, 2024 - Present": [
        "Engineered robust infrastructure for educational platform using Python and ReactJS within an agile organization, implementing secure database queries serving 5,000+ concurrent users",
        "Built automated ETL pipelines with Python that integrate with business management systems, improving data processing efficiency by 35% and enabling seamless data flow",
        "Developed reusable backend services implementing agile methodologies for data protection and user identity management in a product-driven environment",
        "Collaborated with cross-functional teams to design scalable database architecture using PostgreSQL, creating accessible educational resources that improved EEO compliance."
      ]
    },
    {
      "Software Engineer Intern | AlcoveX Product Studio | Jun, 2023 - Nov, 2023": [
        "Designed and developed backend systems using agile methodologies in a product-driven environment, increasing overall system reliability by 40%",
        "Built business intelligence pipelines for analytics processing, implementing AI-powered visualization tools and OOP design patterns for modular code",
        "Created comprehensive test cases for information systems components, ensuring accessibility and EEO compliance in user interfaces",
        "Implemented CI/CD pipelines for automated warehouse management system testing in a collaborative organization, achieving 99.9% deployment success rates."
      ]
    },
    {
      "QA Intern | Medha International": [
        "Developed test automation framework using Python and OOP principles for business management system validation in an agile environment",
        "Built performance testing tools to analyze systems under high load conditions, implementing Slack notifications for real-time alerts",
        "Established streamlined bug reporting workflows to identify software design issues, improving organizational communication",
        "Created data analysis scripts using Python to optimize materials management systems and identify business process improvements."
      ]
    }
  ],
  "projects": [
    {
      "G-Community: Business Networking Platform for College Entrepreneurs": [
        "Architected and developed a reusable backend infrastructure for business management platform supporting 15,000+ active users",
        "Engineered a robust real-time communication system implementing MySQL and database queries for user identity management",
        "Integrated with Slack for notifications and business systems to create seamless information flow between sales and operations."
      ]
    },
    {
      "Dermatological Diagnostics: Deep Learning Framework": [
        "Engineered a CNN-based automated system for classifying nine skin pathologies, creating a distributed architecture that processed 1.2TB of medical images",
        "Implemented distributed training architecture on Amazon Web Services to process large medical datasets, reducing training time by 65%",
        "Achieved 85% diagnostic accuracy with optimized model architecture that reduced inference time by 72% for real-time customer applications",
        "Designed fault-tolerant distributed storage solutions that maintained 99.9% data integrity while meeting HIPAA security requirements."
      ]}
      
      ],
  "skills":[
    "**Programming Languages:** Java, Python, C/C++, JavaScript, R, Golang",
    "**Software Development:** Object-oriented design, Distributed systems, Algorithms, Data structures, Distributed storage",
    "**Cloud & Infrastructure:** Amazon Web Services, Distributed storage, AWS S3, Security implementations, Containerization",
    "**Database Systems:** PostgreSQL, MySQL, Relational databases, Distributed query systems",
    "**Development Tools:** Git, Docker, Jenkins, CI/CD pipelines",
    "**Operating Systems & Working Environment:** Windows, Linux, Ubuntu",
    "**Image Processing:** Computer Vision applications, Image Analysis algorithms, AI-powered imaging solutions",
    "**Additional Skills:** Data mining, Customer-focused UX design, Complexity analysis, Optimization"
    
  ]
  
  }
    data_dev_resume = software_dev_resume
    full_stack_resume = software_dev_resume
    software_testing_resume = software_dev_resume
    job_description = st.text_area("Paste job description here", height=150)
    type_of_role = st.selectbox("Select type of role", ["Software Developer", "Data Science", "Full Stack", "Software Testing"])

    if uploaded_file :
        try:
            # Parse JSON input
            if type_of_role == "Software Developer": resume_data = software_dev_resume
            elif type_of_role == "Data Science": resume_data = data_dev_resume
            elif type_of_role == "Full Stack": resume_data = full_stack_resume
            else : resume_data = software_testing_resume
            
            if st.button("Build Resume"):
                with st.spinner("Processing..."):
                    # Load document
                    doc = load_docx(uploaded_file)
                    
                    # If using OpenAI, get optimized content
                    if job_description :
                        enhanced_data = process_job_description(resume_data, job_description, model_option)
                        if enhanced_data:
                            resume_data = enhanced_data
                            st.success("Resume content optimized for job description!")
                    
                    # Add experience section
                    if "experience" in resume_data and resume_data["experience"]:
                        success = add_content_to_section(doc, "EXPERIENCE", resume_data["experience"])
                        if success:
                            st.success("✅ Added experience section")
                        else:
                            st.error("❌ Failed to add experience section")
                    
                    # Add projects section
                    if "projects" in resume_data and resume_data["projects"]:
                        success = add_content_to_section(doc, "PROJECTS", resume_data["projects"])
                        if success:
                            st.success("✅ Added projects section")
                        else:
                            st.error("❌ Failed to add projects section")
                    
                    if "skills" in resume_data and resume_data["skills"]:
                        success = add_skills_to_section(doc, "TECHNICAL SKILLS", resume_data["skills"])
                        if success:
                            st.success("✅ Added technical skills section")
                        else:
                            st.error("❌ Failed to add technical skills section")
                    
                    # Save and offer download
                    output_docx = save_docx(doc)
                    st.download_button("📥 Download Complete Resume", 
                                      data=output_docx, 
                                      file_name="Complete_Resume.docx",
                                      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except json.JSONDecodeError:
            st.error("Invalid JSON format. Please check your input.")

if __name__ == "__main__":
    main()